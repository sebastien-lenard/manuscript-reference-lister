import csv
import os
import subprocess
from pathlib import Path

import pytest
from docx import Document


@pytest.mark.e2e
def test_docx_file_pipeline_execution(tmp_path: Path) -> None:
    """Validation of full pipeline from a temporary .docx and check structure of CSV
    generated."""
    input_file = tmp_path / "test_manuscript.docx"
    output_csv = tmp_path / "output_test_references.csv"
    doc = Document()

    doc.add_paragraph(
        "Le voila (Lenard et al., 2020), sans parler des erreurs oubliées "
        "(e.g. Fig. 7 in Guns and Vanacker, 2014). Jeu de donnees de Croissant et al. "
        "(2019)."
    )

    doc.add_paragraph("Journals\nGeology\nNature Geoscience\nAnthropocene")

    doc.save(input_file)

    cmd = [
        "uv",
        "run",
        "python",
        "-m",
        "src.manuscript_reference_lister",
        "-v",
        "-f",
        str(input_file),
        "-o",
        str(output_csv),
    ]

    isolated_env = os.environ.copy()
    isolated_env["LOCAL_REPO_DIR_PATH"] = str(tmp_path / "repo")
    isolated_env["LOG_DIR_PATH"] = str(tmp_path / "log")
    isolated_env["OUTPUT_DIR_PATH"] = str(tmp_path / "output")

    result = subprocess.run(
        cmd, capture_output=True, text=True, check=False, env=isolated_env
    )

    assert result.returncode == 0, f"Failed .docx pipeline:\n{result.stderr}"
    assert output_csv.exists(), f"No output file generated: {output_csv}"
    with open(output_csv, encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)

        assert reader.fieldnames == ["Citation", "Status", "Reference"], (
            f"Incorrect columns in output CSV: {reader.fieldnames}"
        )

        rows = list(reader)
        assert len(rows) > 0, "Empty output CSV."

        citations_generes = [row["Citation"] for row in rows]
        assert any("Lenard et al., 2020" in info for info in citations_generes), (
            f"Missing citation 'Lenard et al., 2020'. Found : {citations_generes}"
        )


@pytest.mark.e2e
def test_stdin_pipeline_execution(tmp_path: Path) -> None:
    """Validation of full pipeline from stdin."""

    input_data = "Text (Lenard et al., 2020)\r\nJournals\r\nNature Geoscience"
    output_csv = tmp_path / "stdin_output.csv"
    cmd = ["uv", "run", "references-lister", "-v", "-o", str(output_csv)]

    isolated_env = os.environ.copy()
    isolated_env["LOCAL_REPO_DIR_PATH"] = str(tmp_path / "repo")
    isolated_env["LOG_DIR_PATH"] = str(tmp_path / "log")
    isolated_env["OUTPUT_DIR_PATH"] = str(tmp_path / "output")

    result = subprocess.run(
        cmd,
        input=input_data,
        capture_output=True,
        text=True,
        check=False,
        env=isolated_env,
    )

    assert result.returncode == 0, f"Failed stdin pipeline:\n{result.stderr}"
    assert output_csv.exists(), f"No output file generated: {output_csv}"
    with open(output_csv, encoding="utf-8-sig") as f:
        content = f.read()
        assert "Lenard et al., 2020" in content
